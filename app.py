"""
Job Scraper Dashboard — Flask Backend
Run: python app.py
Open: http://localhost:5000
"""

from flask import Flask, request, jsonify, send_file
import os, json, csv, re, time, threading, random
from pathlib import Path
from collections import Counter
import tempfile

from dotenv import load_dotenv

# If _personal/ exists, load .env and data from there (personal mode).
# If not, use root paths (clean product mode for new users).
BASE_DIR = Path(__file__).parent
PERSONAL_DIR = BASE_DIR / "_personal"
IS_PERSONAL = PERSONAL_DIR.exists()

if IS_PERSONAL:
    load_dotenv(PERSONAL_DIR / ".env")
    DATA_DIR = PERSONAL_DIR / "data"
    OUTPUTS_DIR = PERSONAL_DIR / "outputs"
else:
    load_dotenv(BASE_DIR / ".env")
    DATA_DIR = BASE_DIR / "data"
    OUTPUTS_DIR = BASE_DIR / "outputs"

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

DATA_DIR.mkdir(exist_ok=True)
OUTPUTS_DIR.mkdir(exist_ok=True)

CONFIG_FILE = (PERSONAL_DIR if IS_PERSONAL else BASE_DIR) / "config.json"

JOBS_FILE           = DATA_DIR / "linkedin_jobs.json"
SCORED_FILE         = DATA_DIR / "linkedin_jobs_scored.json"
APPLIED_FILE        = DATA_DIR / "applied.json"
NOT_INTERESTED_FILE = DATA_DIR / "not_interested.json"
REJECTED_FILE       = DATA_DIR / "rejected.json"
INTERESTED_FILE     = DATA_DIR / "interested.json"
ACTION_NEEDED_FILE  = DATA_DIR / "action_needed.json"
EXPIRED_FILE        = DATA_DIR / "expired.json"

# ─── State ────────────────────────────────────────────────────────────────────
state = {
    "running": False,
    "stop_requested": False,
    "stage": "",
    "log": [],
    "jobs": [],
    "scored_jobs": [],
    "clusters": [],
    "study_plan": [],
    "applier_results": [],
    "error": None
}

# Handle to the live pipeline thread so a new Run can force the previous
# session to wind down before starting completely fresh.
pipeline_thread = None

def log(msg):
    # flush=True ensures lines appear in the terminal immediately even when
    # stdout is piped (e.g. under `conda run`, IDE consoles, log files).
    print(msg, flush=True)
    state["log"].append(msg)

# ─── Pacing (anti-rate-limit jitter) ──────────────────────────────────────────
# All delays keep randomness — constant intervals are the easiest bot signal —
# but the tier controls how aggressive the pace is. "safe" is the old overnight
# pace; "balanced" is the interactive default; "fast" is for small/test runs.
SPEED_PROFILES = {
    "fast":     {"between": (2, 4),   "page_ms": (1500, 2500),
                 "batch_every": 0,    "batch_pause": (0, 0),   "login_ms": (3000, 5000)},
    "balanced": {"between": (4, 8),   "page_ms": (2000, 3500),
                 "batch_every": 40,   "batch_pause": (20, 35), "login_ms": (4000, 6000)},
    "safe":     {"between": (8, 15),  "page_ms": (3000, 5000),
                 "batch_every": 25,   "batch_pause": (45, 90), "login_ms": (6000, 9000)},
}

def _profile(config):
    return SPEED_PROFILES.get(config.get("speed", "balanced"), SPEED_PROFILES["balanced"])

def interruptible_sleep(seconds):
    """Sleep in ~0.3s slices, bailing the instant a stop is requested.
    Returns False if interrupted, True if it slept the full duration. This is
    what makes Stop feel instant instead of waiting out a 90s pause."""
    end = time.time() + seconds
    while True:
        remaining = end - time.time()
        if remaining <= 0:
            return True
        if state["stop_requested"]:
            return False
        time.sleep(min(0.3, remaining))

# ─── File Helpers ─────────────────────────────────────────────────────────────
def load_file(path):
    if path.exists():
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_file(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)

def get_timestamp():
    from datetime import datetime
    return datetime.now().strftime("%Y-%m-%d %H:%M")

def remove_from_all(job_id):
    for path in [APPLIED_FILE, NOT_INTERESTED_FILE, REJECTED_FILE,
                 INTERESTED_FILE, ACTION_NEEDED_FILE, EXPIRED_FILE]:
        data = load_file(path)
        if job_id in data:
            data.pop(job_id)
            save_file(path, data)

# ─── Routes ───────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    with open(BASE_DIR / "dashboard.html", encoding="utf-8") as f:
        return f.read()

@app.route("/config")
def get_config():
    """Merge .env secrets (take priority) with config.json preferences."""
    cfg = {}
    # Load preferences from config.json (keywords, location, profession)
    config_path = CONFIG_FILE
    if config_path.exists():
        with open(config_path, encoding="utf-8") as f:
            cfg.update(json.load(f))
    # Secrets from .env always take priority
    for env_key, cfg_key in [("LINKEDIN_EMAIL", "email"),
                              ("LINKEDIN_PASSWORD", "password")]:
        val = os.environ.get(env_key, "")
        if val:
            cfg[cfg_key] = val
    # LLM config — new unified keys with legacy fallback
    cfg.setdefault("llm_provider", os.environ.get("LLM_PROVIDER", "anthropic"))
    cfg.setdefault("llm_model", os.environ.get("LLM_MODEL", "claude-sonnet-4-6"))
    cfg.setdefault("speed", os.environ.get("SPEED", "balanced"))
    api_key = os.environ.get("LLM_API_KEY") or os.environ.get("ANTHROPIC_API_KEY", "")
    if api_key:
        cfg["llm_api_key"] = api_key
    return jsonify(cfg)

def _persist_config(cfg):
    """Save non-sensitive preferences to config.json. Secrets stay in .env."""
    pref_keys = ["keywords", "location", "profession", "llm_provider", "llm_model",
                 "llm_base_url", "speed"]
    out = {}
    # Preserve existing preferences if not in this request
    config_path = CONFIG_FILE
    if config_path.exists():
        try:
            with open(config_path, encoding="utf-8") as f:
                for k in pref_keys:
                    if k in (existing := json.load(f)):
                        out[k] = existing[k]
        except Exception:
            pass
    for k in pref_keys:
        v = cfg.get(k, "")
        if k == "keywords" and isinstance(v, list):
            v = ", ".join(v)
        if v:
            out[k] = v
    try:
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(out, f, indent=2, ensure_ascii=False)
    except Exception as e:
        log(f"Warning: could not save config.json: {e}")

@app.route("/status")
def status():
    return jsonify({
        "running": state["running"],
        "stop_requested": state["stop_requested"],
        "stage": state["stage"],
        "log": state["log"][-50:],
        "job_count": len(state["jobs"]),
        "scored_count": len(state["scored_jobs"]),
        "error": state["error"]
    })

@app.route("/stop", methods=["POST"])
def stop():
    if state["running"]:
        state["stop_requested"] = True
        state["stage"] = "Stopping..."
        log("⚠ Stop requested — finishing current task then stopping...")
        return jsonify({"ok": True})
    return jsonify({"error": "Not running"}), 400

@app.route("/reset", methods=["POST"])
def reset():
    # Always allow reset, even when "running" is True — this rescues the user
    # from a stuck pipeline (e.g. Playwright hang). The background thread (if
    # actually alive) will see stop_requested on its next checkpoint and exit;
    # if it was a stale flag from a crashed run, this just clears it.
    was_running = state["running"]
    state["running"] = False
    state["stop_requested"] = True if was_running else False
    state["stage"] = ""
    state["log"] = []
    state["jobs"] = []
    state["scored_jobs"] = []
    state["clusters"] = []
    state["study_plan"] = []
    state["applier_results"] = []
    state["error"] = None
    if was_running:
        log("⚠ Force-reset while running — any background work will be discarded.")
    else:
        log("Pipeline reset. Ready to run.")
    return jsonify({"ok": True})

@app.route("/run", methods=["POST"])
def run():
    global pipeline_thread

    # Run = a brand-new session. If a previous pipeline is still alive (or left
    # a stale running flag), signal it to stop and let it wind down its browser
    # before we wipe state and start fresh — "like a new app.py was started".
    if state["running"] or (pipeline_thread and pipeline_thread.is_alive()):
        state["stop_requested"] = True
        if pipeline_thread and pipeline_thread.is_alive():
            pipeline_thread.join(timeout=12)
        state["running"] = False

    data = request.form
    files = request.files
    mode = data.get("mode")

    single_run = data.get("single_run", "") in ("1", "true", "True", "on")

    config = {
        "mode": mode,
        "single_run": single_run,
        "speed": "fast" if single_run else (data.get("speed", "") or "balanced"),
        "email": data.get("email", ""),
        "password": data.get("password", ""),
        "xing_email": data.get("xing_email", "") or os.environ.get("XING_EMAIL", ""),
        "xing_password": data.get("xing_password", "") or os.environ.get("XING_PASSWORD", ""),
        "job_sources": data.getlist("job_sources") or ["linkedin"],
        "country": data.get("country", "at"),
        "keywords": [k.strip() for k in data.get("keywords", "").split(",") if k.strip()],
        "location": data.get("location", ""),
        "llm_provider": data.get("llm_provider", "") or os.environ.get("LLM_PROVIDER", "anthropic"),
        "llm_model": data.get("llm_model", "") or os.environ.get("LLM_MODEL", "claude-sonnet-4-6"),
        "llm_api_key": data.get("llm_api_key", "") or os.environ.get("LLM_API_KEY", "") or os.environ.get("ANTHROPIC_API_KEY", ""),
        "llm_base_url": data.get("llm_base_url", "") or os.environ.get("LLM_BASE_URL", ""),
        "time_filter": data.get("time_filter", "r604800"),
        "pages": 1 if single_run else int(data.get("pages", 5)),
        "profession": data.get("profession", "").strip(),
        "resume_text": ""
    }

    # A single test run is login → 1 JD → score. Always score, never cluster
    # (clustering one job is meaningless), regardless of the selected mode.
    if single_run:
        config["mode"] = mode = "with_scoring"

    if "resume" in files:
        resume_file = files["resume"]
        suffix = Path(resume_file.filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            resume_file.save(tmp.name)
            try:
                if suffix == ".pdf":
                    import pdfplumber
                    with pdfplumber.open(tmp.name) as pdf:
                        config["resume_text"] = "\n".join(p.extract_text() or "" for p in pdf.pages)
                elif suffix in [".docx", ".doc"]:
                    import docx
                    doc = docx.Document(tmp.name)
                    config["resume_text"] = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                log(f"Warning: could not extract resume text: {e}")

    # Persist form values so they prefill on next page load
    _persist_config(config)

    # Validate API key before launching pipeline (fail fast)
    if mode in ("with_scoring", "full"):
        try:
            _llm_ping(config)
        except Exception as e:
            return jsonify({"error": f"LLM API key / connection invalid: {e}"}), 400

    state["running"] = True
    state["stop_requested"] = False
    state["log"] = []
    state["jobs"] = []
    state["scored_jobs"] = []
    state["clusters"] = []
    state["study_plan"] = []
    state["error"] = None
    state["stage"] = "Starting..."

    thread = threading.Thread(target=run_pipeline, args=(config,))
    thread.daemon = True
    thread.start()
    pipeline_thread = thread

    return jsonify({"ok": True})

@app.route("/download/<filetype>")
def download(filetype):
    if filetype == "csv":
        from datetime import datetime
        date_str = datetime.now().strftime("%Y-%m-%d")
        path = generate_csv()
        return send_file(path, as_attachment=True, download_name=f"job_results_{date_str}.csv")
    elif filetype == "skills":
        return send_file(generate_skills_txt(), as_attachment=True, download_name="skill_clusters.txt")
    elif filetype == "plan":
        return send_file(generate_plan_txt(), as_attachment=True, download_name="study_plan.txt")
    return "Not found", 404

# ─── Job Status Routes ────────────────────────────────────────────────────────
@app.route("/applied", methods=["GET"])
def get_applied():
    return jsonify(load_file(APPLIED_FILE))

@app.route("/applied/mark", methods=["POST"])
def mark_applied():
    data = request.get_json()
    job = data.get("job")
    if not job:
        return jsonify({"error": "No job provided"}), 400
    job_id = job.get("url") or job.get("title")
    remove_from_all(job_id)
    job["applied_at"] = get_timestamp()
    applied = load_file(APPLIED_FILE)
    applied[job_id] = job
    save_file(APPLIED_FILE, applied)
    return jsonify({"ok": True})

@app.route("/applied/unmark", methods=["POST"])
def unmark_applied():
    data = request.get_json()
    job_id = data.get("job_id")
    applied = load_file(APPLIED_FILE)
    applied.pop(job_id, None)
    save_file(APPLIED_FILE, applied)
    return jsonify({"ok": True})

@app.route("/applied/reject", methods=["POST"])
def reject_applied():
    data = request.get_json()
    job_id = data.get("job_id")
    applied = load_file(APPLIED_FILE)
    job = applied.pop(job_id, None)
    save_file(APPLIED_FILE, applied)
    if job:
        job["rejected_at"] = get_timestamp()
        rejected = load_file(REJECTED_FILE)
        rejected[job_id] = job
        save_file(REJECTED_FILE, rejected)
    return jsonify({"ok": True})

@app.route("/applied/switch_type", methods=["POST"])
def switch_application_type():
    data = request.get_json()
    job_id = data.get("job_id")
    new_type = data.get("application_type")
    applied = load_file(APPLIED_FILE)
    if job_id in applied:
        applied[job_id]["application_type"] = new_type
        save_file(APPLIED_FILE, applied)
    return jsonify({"ok": True})

@app.route("/not_interested", methods=["GET"])
def get_not_interested():
    return jsonify(load_file(NOT_INTERESTED_FILE))

@app.route("/not_interested/mark", methods=["POST"])
def mark_not_interested():
    data = request.get_json()
    job = data.get("job")
    if not job:
        return jsonify({"error": "No job provided"}), 400
    job_id = job.get("url") or job.get("title")
    remove_from_all(job_id)
    job["hidden_at"] = get_timestamp()
    ni = load_file(NOT_INTERESTED_FILE)
    ni[job_id] = job
    save_file(NOT_INTERESTED_FILE, ni)
    return jsonify({"ok": True})

@app.route("/not_interested/restore", methods=["POST"])
def restore_not_interested():
    data = request.get_json()
    job_id = data.get("job_id")
    ni = load_file(NOT_INTERESTED_FILE)
    ni.pop(job_id, None)
    save_file(NOT_INTERESTED_FILE, ni)
    return jsonify({"ok": True})

@app.route("/rejected", methods=["GET"])
def get_rejected():
    return jsonify(load_file(REJECTED_FILE))

@app.route("/rejected/restore", methods=["POST"])
def restore_rejected():
    data = request.get_json()
    job_id = data.get("job_id")
    rejected = load_file(REJECTED_FILE)
    job = rejected.pop(job_id, None)
    save_file(REJECTED_FILE, rejected)
    if job:
        job.pop("rejected_at", None)
        job["applied_at"] = job.get("applied_at", get_timestamp())
        applied = load_file(APPLIED_FILE)
        applied[job_id] = job
        save_file(APPLIED_FILE, applied)
    return jsonify({"ok": True})

@app.route("/interested", methods=["GET"])
def get_interested():
    return jsonify(load_file(INTERESTED_FILE))

@app.route("/interested/mark", methods=["POST"])
def mark_interested():
    data = request.get_json()
    job = data.get("job")
    if not job:
        return jsonify({"error": "No job provided"}), 400
    job_id = job.get("url") or job.get("title")
    remove_from_all(job_id)
    job["interested_at"] = get_timestamp()
    interested = load_file(INTERESTED_FILE)
    interested[job_id] = job
    save_file(INTERESTED_FILE, interested)
    return jsonify({"ok": True})

@app.route("/action_needed", methods=["GET"])
def get_action_needed():
    applied = load_file(APPLIED_FILE)
    action_needed = {k: v for k, v in applied.items() if v.get("action_needed_at")}
    return jsonify(action_needed)

@app.route("/action_needed/mark", methods=["POST"])
def mark_action_needed():
    data = request.get_json()
    job_id = data.get("job_id")
    applied = load_file(APPLIED_FILE)
    if job_id in applied:
        applied[job_id]["action_needed_at"] = get_timestamp()
        save_file(APPLIED_FILE, applied)
    return jsonify({"ok": True})

@app.route("/expired", methods=["GET"])
def get_expired():
    # The client only needs the set of expired job ids (to grey out expired jobs),
    # not the full records — returning every expired job dict was a 28MB payload
    # that stalled page loads on mobile. Return just the keys.
    return jsonify(list(load_file(EXPIRED_FILE).keys()))

@app.route("/expired/mark", methods=["POST"])
def mark_expired():
    data = request.get_json()
    job = data.get("job")
    if not job:
        return jsonify({"error": "No job provided"}), 400
    job_id = job.get("url") or job.get("title")
    remove_from_all(job_id)
    job["expired_at"] = get_timestamp()
    expired = load_file(EXPIRED_FILE)
    expired[job_id] = job
    save_file(EXPIRED_FILE, expired)
    return jsonify({"ok": True})

@app.route("/salary_stats")
def salary_stats():
    jobs = state["scored_jobs"] or state["jobs"]
    annual  = [j for j in jobs if classify_salary(j.get("salary","")) == "annual"]
    hourly  = [j for j in jobs if classify_salary(j.get("salary","")) == "hourly"]
    missing = [j for j in jobs if classify_salary(j.get("salary","")) == "missing"]

    def summarize(group):
        return [{
            "title": j.get("title",""),
            "company": j.get("company",""),
            "location": j.get("location",""),
            "salary": j.get("salary",""),
            "source": j.get("source",""),
            "fit_score": j.get("fit_score",""),
            "response_probability": j.get("response_probability",""),
            "matched_skills": j.get("matched_skills", []),
            "missing_skills": j.get("missing_skills", []),
            "verdict": j.get("verdict",""),
            "url": j.get("url","")
        } for j in sorted(group, key=lambda x: x.get("fit_score",0)
                          if x.get("fit_score","") != "" else 0, reverse=True)]

    return jsonify({"annual": summarize(annual), "hourly": summarize(hourly), "missing": summarize(missing)})

@app.route("/load_csv", methods=["POST"])
def load_csv():
    if "csv_file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["csv_file"]
    if not file.filename.endswith(".csv"):
        return jsonify({"error": "File must be a .csv"}), 400
    try:
        import io
        content_str = file.read().decode("utf-8")
        reader = csv.reader(io.StringIO(content_str))
        headers = next(reader, None)
        if not headers:
            return jsonify({"error": "Empty CSV"}), 400

        def col(name):
            for i, h in enumerate(headers):
                if h.lower().strip() == name.lower():
                    return i
            return None

        idx = {
            "title": col("title"), "company": col("company"), "location": col("location"),
            "salary": col("salary"), "salary_type": col("salary type"), "source": col("source"),
            "fit_score": col("fit score"), "response_probability": col("response probability"),
            "missing_skills": col("missing skills"), "verdict": col("verdict"), "url": col("url"),
        }

        annual, hourly, missing_salary = [], [], []

        for row in reader:
            if not row or not any(row) or row[0].startswith("---"):
                continue

            def get(key):
                i = idx.get(key)
                return row[i].strip() if i is not None and i < len(row) else ""

            salary_type = get("salary_type") or classify_salary(get("salary"))
            missing = [s.strip() for s in get("missing_skills").split("|") if s.strip()]

            job = {
                "title": get("title"), "company": get("company"), "location": get("location"),
                "salary": get("salary"), "source": get("source"), "fit_score": get("fit_score"),
                "response_probability": get("response_probability"),
                "missing_skills": missing, "matched_skills": [],
                "verdict": get("verdict"), "url": get("url"),
            }

            if salary_type == "annual": annual.append(job)
            elif salary_type == "hourly": hourly.append(job)
            else: missing_salary.append(job)

        def sort_by_fit(group):
            def fit_key(j):
                try: return int(j.get("fit_score", 0))
                except: return 0
            return sorted(group, key=fit_key, reverse=True)

        return jsonify({
            "annual": sort_by_fit(annual), "hourly": sort_by_fit(hourly),
            "missing": sort_by_fit(missing_salary),
            "total": len(annual) + len(hourly) + len(missing_salary)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ─── Provider / Model Registry ────────────────────────────────────────────────
# Verified June 2026. For each provider:
#   sdk       — which client path ("anthropic" native, or "openai"-compatible)
#   base_url  — OpenAI-compatible endpoint (None = SDK default / native host)
#   models    — {model_id: max_output_tokens}. NOTE: context window (how much
#               INPUT a model accepts, often 200K–1M) is NOT the same as the
#               max OUTPUT tokens a single call may return. The number below is
#               the OUTPUT ceiling — that's what max_tokens must be clamped to.
PROVIDERS = {
    "anthropic": {
        "sdk": "anthropic",
        "base_url": None,
        "models": {
            "claude-opus-4-8":   64000,
            "claude-sonnet-4-6": 64000,
            "claude-haiku-4-5":  32000,
        },
    },
    "deepseek": {
        "sdk": "openai",
        "base_url": "https://api.deepseek.com",
        "models": {
            "deepseek-chat":     65536,
            "deepseek-reasoner": 65536,
            "deepseek-v4-pro":   65536,
            "deepseek-v4-flash": 65536,
        },
    },
    "gemini": {
        "sdk": "openai",
        "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
        "models": {
            "gemini-2.5-pro":          65536,
            "gemini-2.5-flash":        65536,
            "gemini-3-flash-preview":  65536,
        },
    },
    "openai": {
        "sdk": "openai",
        "base_url": None,
        "models": {
            "gpt-4o":      16000,
            "gpt-4o-mini": 16000,
        },
    },
    "llama": {
        "sdk": "openai",
        "base_url": "https://api.llama.com/compat/v1/",
        "models": {
            "Llama-4-Maverick-17B-128E-Instruct-FP8": 16000,
            "Llama-4-Scout-17B-16E-Instruct":         16000,
        },
    },
    "custom": {
        "sdk": "openai",
        "base_url": None,   # user supplies llm_base_url
        "models": {},        # user types any model id
    },
}

# Fallback output ceiling for unknown / user-typed models.
DEFAULT_MAX_OUTPUT = 8000

# Automatic parallel-scoring concurrency per provider. Tuned to stay comfortably
# under each provider's default rate limits — the user never sets this. DeepSeek
# is generous; the rest are kept conservative so low API tiers don't hit 429s.
PROVIDER_WORKERS = {
    "anthropic": 4,
    "deepseek":  8,
    "gemini":    4,
    "openai":    4,
    "llama":     3,
    "custom":    3,
}

def _auto_workers(config, job_count):
    """Pick a safe number of parallel scoring workers automatically.
    Never exceeds the number of jobs, and an optional SCORE_WORKERS env var
    lets a power user override without any UI."""
    if config.get("single_run") or job_count <= 1:
        return 1
    env = os.environ.get("SCORE_WORKERS")
    if env:
        try:
            return max(1, min(int(env), 12))
        except ValueError:
            pass
    base = PROVIDER_WORKERS.get(config.get("llm_provider", "anthropic"), 3)
    return max(1, min(base, job_count))

# Substrings that mean the browser/context is gone for good. Once we see one,
# retrying just burns 45s per attempt for hours (the original 4-hour hang) —
# so we abort the scrape and keep whatever we already have.
_FATAL_PW_SUBSTRINGS = (
    "target page, context or browser has been closed",
    "browser has been closed",
    "context has been closed",
    "context was destroyed",
    "browser closed",
    "connection closed",
    "has crashed",
    "crashed",
)

def _is_fatal_pw_error(e):
    msg = str(e).lower()
    return any(s in msg for s in _FATAL_PW_SUBSTRINGS)

def _provider_spec(provider):
    # Unknown provider name → treat as a custom OpenAI-compatible endpoint.
    return PROVIDERS.get(provider, PROVIDERS["custom"])

def _model_max_output(provider, model):
    return _provider_spec(provider).get("models", {}).get(model, DEFAULT_MAX_OUTPUT)

def _resolve_max_tokens(config, want):
    """Clamp a desired output-token budget to the chosen model's real ceiling.
    Fixes the old hardcoded 2000 cap that truncated output on large models."""
    cap = _model_max_output(config.get("llm_provider", "anthropic"),
                            config.get("llm_model", ""))
    return max(256, min(want, cap))

# ─── LLM Helpers (multi-provider) ─────────────────────────────────────────────
def _llm_client(config):
    """Return (sdk_name, client) for the configured provider."""
    provider = config.get("llm_provider", "anthropic")
    api_key = config.get("llm_api_key", "")
    spec = _provider_spec(provider)
    # Explicit user override (custom) wins over the registry default.
    url = config.get("llm_base_url", "") or spec.get("base_url")

    if spec["sdk"] == "anthropic":
        import anthropic
        kwargs = {"api_key": api_key}
        if url:
            kwargs["base_url"] = url
        return ("anthropic", anthropic.Anthropic(**kwargs))
    else:
        import openai
        kwargs = {"api_key": api_key}
        if url:
            kwargs["base_url"] = url
        return ("openai", openai.OpenAI(**kwargs))

def _llm_ping(config):
    """Quick validation call to check API key / connectivity."""
    sdk, client = _llm_client(config)
    model = config.get("llm_model", "claude-sonnet-4-6")
    if sdk == "anthropic":
        client.messages.create(model=model, max_tokens=1,
                               messages=[{"role": "user", "content": "ping"}])
    else:
        client.chat.completions.create(model=model, max_tokens=1,
                                       messages=[{"role": "user", "content": "ping"}])

def _llm_chat(config, prompt, max_tokens=2000, json_mode=False):
    """Send a prompt and return the text response. Routes to correct SDK.
    Output budget is clamped to the model's ceiling and usage is logged so the
    real input/output token counts are visible in the live log.

    json_mode forces the OpenAI-compatible path to emit a JSON object
    (response_format), which prevents structurally invalid JSON at the source.
    Endpoints that don't support response_format fall back to a plain call."""
    sdk, client = _llm_client(config)
    model = config.get("llm_model", "claude-sonnet-4-6")
    capped = _resolve_max_tokens(config, max_tokens)
    if sdk == "anthropic":
        resp = client.messages.create(
            model=model, max_tokens=capped,
            messages=[{"role": "user", "content": prompt}])
        try:
            u = resp.usage
            log(f"  · tokens in={u.input_tokens} out={u.output_tokens} (cap {capped})")
        except Exception:
            pass
        return (resp.content[0].text or "").strip()
    else:
        kwargs = {"model": model, "max_tokens": capped,
                  "messages": [{"role": "user", "content": prompt}]}
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        try:
            resp = client.chat.completions.create(**kwargs)
        except Exception:
            # Some custom OpenAI-compatible endpoints reject response_format —
            # retry once without it rather than failing the whole call.
            if json_mode:
                kwargs.pop("response_format", None)
                resp = client.chat.completions.create(**kwargs)
            else:
                raise
        try:
            u = resp.usage
            log(f"  · tokens in={u.prompt_tokens} out={u.completion_tokens} (cap {capped})")
        except Exception:
            pass
        return (resp.choices[0].message.content or "").strip()

# ─── Pipeline ─────────────────────────────────────────────────────────────────
def run_pipeline(config):
    single = config.get("single_run", False)
    try:
        mode = config["mode"]
        sources = config.get("job_sources", ["linkedin"])

        if single:
            log(f"▶ Test run: one job per ticked source ({', '.join(sources)}) — "
                f"scrape → score.")
        state["stage"] = ("Test run: scraping one job per source..." if single
                          else f"Scraping {', '.join(s.title() for s in sources)}...")

        # Run every ticked source. On a Test Run each scraper stops at its first
        # usable job description, so we get one job PER source (not just LinkedIn).
        def _src_count(src):
            return len([j for j in state["jobs"]
                        if j.get("source", "linkedin") == src])

        scraper_map = {
            "linkedin": scrape_jobs,
            "indeed":   scrape_indeed_jobs,
            "karriere": scrape_karriere_jobs,
            "xing":     scrape_xing_jobs,
        }
        for src in ("linkedin", "indeed", "karriere", "xing"):
            if src not in sources or state["stop_requested"]:
                continue
            before = _src_count(src)
            try:
                scraper_map[src](config)
            except Exception as e:
                if _is_fatal_pw_error(e):
                    log(f"  ✗ {src}: browser session lost — skipping this source. "
                        f"({str(e)[:80]})")
                else:
                    log(f"  ✗ {src}: scrape error — skipping. ({str(e)[:120]})")
            got = _src_count(src) - before
            with_desc = len([j for j in state["jobs"]
                             if j.get("source", "linkedin") == src
                             and j.get("description")])
            mark = "✓" if with_desc else "✗"
            log(f"  {mark} {src}: {got} job(s), {with_desc} with description"
                + (" [test]" if single else ""))

        if state["stop_requested"]:
            raise StopIteration("Stopped by user after scraping")

        # Single run must end with at least one usable JD somewhere; surface a
        # clear error only if EVERY source came back empty.
        if single and not any(j.get("description") for j in state["jobs"]):
            raise RuntimeError("No source produced a job description. Check the log "
                               "for per-source failures (login, blocks, selectors).")

        if mode in ["with_scoring", "full"]:
            state["stage"] = "Scoring jobs with AI..."
            score_jobs(config)

        if mode == "full" and not state["stop_requested"]:
            state["stage"] = "Generating skill clusters and study plan..."
            generate_clusters_and_plan(config)

        # A test run is throwaway — don't pollute the persistent job database.
        if not single:
            save_jobs()
        state["stage"] = "Done ✓"
        state["running"] = False
        if single:
            n = len([j for j in state["jobs"] if j.get("description")])
            srcs = sorted({j.get("source", "linkedin") for j in state["jobs"]
                           if j.get("description")})
            log(f"✓ Test run complete — {n} job(s) scored across "
                f"{len(srcs)} source(s): {', '.join(srcs) or 'none'}.")
        else:
            log("✓ Pipeline complete. Download your files below.")

    except StopIteration as e:
        save_jobs()
        state["stage"] = "Stopped ⚠"
        state["running"] = False
        state["stop_requested"] = False
        log(f"⚠ Pipeline stopped: {e}")
        log("Partial results saved. Download CSV or reset to start fresh.")
    except Exception as e:
        import traceback
        state["error"] = str(e)
        state["stage"] = "Error"
        state["running"] = False
        log(f"✗ Error: {e}")
        log(traceback.format_exc())

def save_jobs():
    existing = {}
    if JOBS_FILE.exists():
        with open(JOBS_FILE, encoding="utf-8") as f:
            for j in json.load(f):
                existing[j["id"]] = j
    for job in state["jobs"]:
        existing[job["id"]] = job
    with open(JOBS_FILE, "w", encoding="utf-8") as f:
        json.dump(list(existing.values()), f, indent=2)

    if state["scored_jobs"]:
        scored_existing = {}
        if SCORED_FILE.exists():
            with open(SCORED_FILE, encoding="utf-8") as f:
                for j in json.load(f):
                    scored_existing[j["id"]] = j
        for job in state["scored_jobs"]:
            scored_existing[job["id"]] = job
        with open(SCORED_FILE, "w", encoding="utf-8") as f:
            json.dump(list(scored_existing.values()), f, indent=2)

    log(f"Saved {len(existing)} jobs to data/")

# ─── Helper: extract description from page body ───────────────────────────────
def extract_description(page):
    """Extract job description. LinkedIn uses hashed class names so we rely
    on text landmarks rather than CSS selectors."""
    # Landmarks LinkedIn uses for the JD section (English + German)
    landmarks = ["About the job", "Job description",
                 "Über den Job", "Stellenbeschreibung"]

    # Find which landmark is on the page
    try:
        body = page.inner_text("body")
    except:
        return ""
    marker = None
    for lm in landmarks:
        if lm in body:
            marker = lm
            break
    if not marker:
        return ""

    # Try section/article containing the landmark for cleaner extraction
    for tag in ["section", "article"]:
        try:
            container = page.locator(tag).filter(has_text=marker).first
            if container.count():
                text = container.inner_text()
                idx = text.find(marker)
                if idx >= 0:
                    desc = text[idx + len(marker):].strip()
                    if len(desc) > 200:
                        return desc[:12000]
        except:
            pass

    # Fallback: body text split on the landmark
    desc = body.split(marker)[-1].strip()
    return desc[:12000] if desc else ""

# ─── Helper: extract salary from page body ────────────────────────────────────
def extract_salary(page):
    """Returns (confidence, salary_text). confidence: 'verified' | 'unverified' | 'none'."""
    salary_keywords = [
        "salary", "gehalt", "compensation", "vergütung", "€", "$", "£",
        "/yr", "/hr", "per year", "per hour", "annually", "jährlich",
        "monatlich", "hourly", "k/y", "pro jahr", "pro monat",
    ]

    def _has_salary_keyword(line):
        lower = line.lower()
        return any(kw.lower() in lower for kw in salary_keywords)

    def _clean(text):
        return re.sub(r'\s*·.*$', '', text).strip()

    salary_patterns = [
        r'[\$€£]\s*[\d,\.]+\s*[kK]?\s*[-–]\s*[\$€£]?\s*[\d,\.]+\s*[kK]?',
        r'[\d,\.]+\s*[kK]?\s*[-–]\s*[\d,\.]+\s*[kK]?\s*(EUR|USD|GBP|€|\$)',
        r'[\$€£]\s*[\d,\.]+\s*(per hour|per year|\/hr|\/yr|annually)',
    ]

    try:
        # High confidence: LinkedIn's own salary elements
        for sel in [".salary-compensation", "[class*='salary']",
                     ".jobs-unified-top-card__job-insight"]:
            try:
                el = page.query_selector(sel)
                if el:
                    text = el.inner_text().strip()
                    if text and 3 < len(text) < 200 and _has_salary_keyword(text):
                        return ("verified", _clean(text))
            except:
                pass

        # Medium confidence: regex match on a line with a salary keyword nearby
        body = page.inner_text("body")
        for line in body.split("\n"):
            line = line.strip()
            if not _has_salary_keyword(line):
                continue
            for pattern in salary_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    return ("unverified", _clean(line))
    except:
        pass
    return ("none", "")

# ─── Step 1a: Indeed Scraper (Playwright, no login needed) ──────────────────
def scrape_indeed_jobs(config):
    """Scrape jobs from Indeed using Playwright. No login required.
    Probe-confirmed: job cards are in the DOM regardless of cookie banner state.
    Cookie banner is dismissed via JS click on first page only."""
    from playwright.sync_api import sync_playwright

    country = config.get("country", "at")
    domain_map = {"at": "at.indeed.com", "com": "www.indeed.com", "de": "de.indeed.com"}
    base_url = f"https://{domain_map.get(country, 'www.indeed.com')}"

    tf = config.get("time_filter", "r604800")
    fromage = {"r86400": "1", "r604800": "7", "r2592000": "30"}.get(tf, "7")

    location = config.get("location", "")
    pages = config.get("pages", 5)
    prof = _profile(config)
    single = config.get("single_run", False)

    seen_ids = {j["id"] for j in state["jobs"]}
    new_count = 0

    _EXTRACT_JS = """
    () => {
        const results = [];
        for (const card of document.querySelectorAll('div.job_seen_beacon')) {
            const link = card.querySelector('a[data-jk]');
            const jk   = link ? link.getAttribute('data-jk') : '';
            const titleEl   = card.querySelector('h2.jobTitle span[title], h2.jobTitle span');
            const companyEl = card.querySelector('[data-testid="company-name"], .companyName');
            const locEl     = card.querySelector('[data-testid="text-location"], .companyLocation');
            const salaryEl  = card.querySelector('.salary-snippet-container, .attribute_snippet');
            const title   = titleEl   ? titleEl.textContent.trim()   : '';
            const company = companyEl ? companyEl.textContent.trim() : '';
            const loc     = locEl     ? locEl.textContent.trim()     : '';
            const salary  = salaryEl  ? salaryEl.textContent.trim()  : '';
            const href    = link ? link.href : '';
            if (title && (jk || href)) {
                results.push({ jk, title, company, loc, salary, href });
            }
        }
        return results;
    }
    """

    _COOKIE_JS = """
    () => {
        const targets = ['Alle Cookies akzeptieren', 'Accept all cookies',
                         'Alle ablehnen', 'Reject all'];
        for (const btn of document.querySelectorAll('button')) {
            const t = btn.textContent.trim();
            if (targets.some(x => t.includes(x))) { btn.click(); return true; }
        }
        return false;
    }
    """

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            ),
            viewport={"width": 1280, "height": 900},
        )
        pg = context.new_page()
        cookie_done = [False]   # list so it's mutable from inner scope

        for keyword in config["keywords"]:
            if state["stop_requested"]:
                break
            log(f"Indeed: scraping '{keyword}'...")

            for page_num in range(pages):
                if state["stop_requested"]:
                    break
                start = page_num * 10
                search_url = (
                    f"{base_url}/jobs?q={keyword.replace(' ', '+')}"
                    f"&l={location.replace(' ', '+')}"
                    f"&fromage={fromage}&start={start}&sort=date"
                )
                try:
                    pg.goto(search_url, wait_until="domcontentloaded", timeout=30000)
                    # Wait for Indeed to inject job cards (probe confirmed ~3s needed)
                    try:
                        pg.wait_for_selector("div.job_seen_beacon", timeout=8000)
                    except Exception:
                        pg.wait_for_timeout(4000)  # fallback flat wait
                except Exception as e:
                    log(f"  Indeed: nav error page {page_num+1}: {e}")
                    break

                # Probe confirmed: cards ARE in DOM even with banner visible.
                jobs_data = pg.evaluate(_EXTRACT_JS)

                # If no cards found, try JS cookie dismiss and re-extract once.
                if not jobs_data and not cookie_done[0]:
                    clicked = pg.evaluate(_COOKIE_JS)
                    if clicked:
                        cookie_done[0] = True
                        pg.wait_for_timeout(1500)
                        jobs_data = pg.evaluate(_EXTRACT_JS)

                # Dismiss banner on first successful page (cleanup, no re-extract needed)
                if not cookie_done[0]:
                    pg.evaluate(_COOKIE_JS)
                    cookie_done[0] = True

                log(f"  Indeed page {page_num+1}: {len(jobs_data)} cards")
                if not jobs_data:
                    snippet = pg.evaluate("() => document.body.innerText.substring(0, 120)")
                    log(f"  Indeed: page snippet: {snippet}")
                    break

                for jd in jobs_data:
                    if state["stop_requested"]:
                        break
                    jk = jd.get("jk", "")
                    job_id = f"indeed_{jk}" if jk else f"indeed_{abs(hash(jd.get('href', '')))}"
                    if job_id in seen_ids or not jd["title"]:
                        continue
                    seen_ids.add(job_id)
                    url = f"{base_url}/viewjob?jk={jk}" if jk else jd.get("href", "")
                    state["jobs"].append({
                        "id": job_id,
                        "title": jd["title"],
                        "company": jd["company"],
                        "location": jd["loc"] or location,
                        "salary": jd["salary"],
                        "url": url,
                        "source": "indeed",
                        "description": "",
                        "easy_apply": False,
                    })
                    new_count += 1

                delay = random.uniform(*prof["between"])
                if not interruptible_sleep(delay):
                    break

                # Test run: one page of cards is enough to find a sample job.
                if single:
                    break

            delay = random.uniform(*prof["between"])
            if not interruptible_sleep(delay):
                break

            # Test run: only the first keyword.
            if single:
                break

        # ── Fetch descriptions ────────────────────────────────────────────────
        indeed_new = [j for j in state["jobs"]
                      if j.get("source") == "indeed" and not j.get("description")]
        if indeed_new:
            log(f"Indeed: fetching descriptions for {len(indeed_new)} jobs"
                + (" (test: stop at first)..." if single else "..."))
            kept = None
            for idx, job in enumerate(indeed_new):
                if state["stop_requested"]:
                    break
                try:
                    pg.goto(job["url"], wait_until="domcontentloaded", timeout=25000)
                    pg.wait_for_timeout(1500)
                    pg.evaluate(_COOKIE_JS)   # dismiss if banner re-appears
                    desc = pg.evaluate("""
                    () => {
                        const el = document.querySelector(
                            '#jobDescriptionText, .jobsearch-jobDescriptionText'
                        );
                        return el ? el.innerText.trim() : '';
                    }
                    """)
                    if desc:
                        job["description"] = desc[:12000]
                except Exception as e:
                    log(f"  Indeed: desc error for {job['title']}: {e}")
                if single:
                    if job.get("description"):
                        kept = job
                        log("  ✓ Indeed: got one description — keeping it for test run.")
                        break
                    if idx + 1 >= 4:
                        log("  ⚠ Indeed: no description in first 4 cards.")
                        break
                    continue
                delay = random.uniform(1.5, 3.0)
                if not interruptible_sleep(delay):
                    break

            # Test run: drop every Indeed card except the one sample we scored.
            if single:
                state["jobs"] = [j for j in state["jobs"]
                                 if j.get("source") != "indeed" or j is kept]

        try:
            context.close()
        except Exception:
            pass
        browser.close()

    log(f"✓ Indeed: {new_count} new jobs collected.")


def scrape_karriere_jobs(config):
    """Scrape jobs from karriere.at (Austria's largest board). No login required.
    Server-rendered HTML — cards live in `li.m-jobsList__item` on the search page,
    detail text in `.m-jobContent`. Austria-only; the country setting is ignored."""
    from playwright.sync_api import sync_playwright

    if config.get("country", "at") != "at":
        log("  karriere.at is Austria-only — running it regardless of country setting.")

    location = config.get("location", "")
    pages = config.get("pages", 5)
    prof = _profile(config)
    single = config.get("single_run", False)

    seen_ids = {j["id"] for j in state["jobs"]}
    new_count = 0

    def _slug(kw):
        import re as _re
        s = kw.lower().strip()
        s = s.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
        s = _re.sub(r"[^a-z0-9]+", "-", s).strip("-")
        return s or "jobs"

    _EXTRACT_JS = """
    () => {
        const results = [];
        for (const card of document.querySelectorAll('li.m-jobsList__item')) {
            const link = card.querySelector('a.m-jobsListItem__titleLink');
            if (!link) continue;
            const href = link.href || '';
            const title = link.textContent.trim();
            const companyEl = card.querySelector('.m-jobsListItem__companyName, .m-jobsListItem__company');
            const locEl = card.querySelector('.m-jobsListItem__location');
            const company = companyEl ? companyEl.textContent.trim() : '';
            const loc = locEl ? locEl.textContent.trim() : '';
            if (title && href) results.push({ href, title, company, loc });
        }
        return results;
    }
    """

    _COOKIE_JS = """
    () => {
        const targets = ['Alle akzeptieren', 'Akzeptieren', 'Accept all', 'Zustimmen'];
        for (const btn of document.querySelectorAll('button')) {
            const t = btn.textContent.trim();
            if (targets.some(x => t.includes(x))) { btn.click(); return true; }
        }
        return false;
    }
    """

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            ),
            viewport={"width": 1280, "height": 900},
            locale="de-AT",
        )
        pg = context.new_page()
        cookie_done = [False]

        for keyword in config["keywords"]:
            if state["stop_requested"]:
                break
            log(f"karriere.at: scraping '{keyword}'...")

            for page_num in range(pages):
                if state["stop_requested"]:
                    break
                loc_param = f"&locations={location.replace(' ', '+')}" if location else ""
                search_url = (
                    f"https://www.karriere.at/jobs/{_slug(keyword)}"
                    f"?page={page_num + 1}{loc_param}"
                )
                try:
                    pg.goto(search_url, wait_until="domcontentloaded", timeout=30000)
                    try:
                        pg.wait_for_selector("li.m-jobsList__item", timeout=8000)
                    except Exception:
                        pg.wait_for_timeout(3000)
                except Exception as e:
                    log(f"  karriere.at: nav error page {page_num+1}: {e}")
                    break

                if not cookie_done[0]:
                    pg.evaluate(_COOKIE_JS)
                    cookie_done[0] = True
                    pg.wait_for_timeout(800)

                jobs_data = pg.evaluate(_EXTRACT_JS)
                log(f"  karriere.at page {page_num+1}: {len(jobs_data)} cards")
                if not jobs_data:
                    break

                for jd in jobs_data:
                    if state["stop_requested"]:
                        break
                    import re as _re
                    m = _re.search(r"/jobs/(\d+)", jd["href"])
                    job_id = f"karriere_{m.group(1)}" if m else f"karriere_{abs(hash(jd['href']))}"
                    if job_id in seen_ids or not jd["title"]:
                        continue
                    seen_ids.add(job_id)
                    state["jobs"].append({
                        "id": job_id,
                        "title": jd["title"],
                        "company": jd["company"],
                        "location": jd["loc"] or location,
                        "salary": "",
                        "url": jd["href"],
                        "source": "karriere",
                        "description": "",
                        "easy_apply": False,
                    })
                    new_count += 1

                if not interruptible_sleep(random.uniform(*prof["between"])):
                    break
                if single:
                    break

            if not interruptible_sleep(random.uniform(*prof["between"])):
                break
            if single:
                break

        # ── Fetch descriptions ────────────────────────────────────────────────
        kar_new = [j for j in state["jobs"]
                   if j.get("source") == "karriere" and not j.get("description")]
        if kar_new:
            log(f"karriere.at: fetching descriptions for {len(kar_new)} jobs"
                + (" (test: stop at first)..." if single else "..."))
            kept = None
            for idx, job in enumerate(kar_new):
                if state["stop_requested"]:
                    break
                try:
                    pg.goto(job["url"], wait_until="domcontentloaded", timeout=25000)
                    pg.wait_for_timeout(1500)
                    body = pg.evaluate("() => document.body.innerText.toLowerCase()")
                    # Skip expired listings — same marker used by the staleness cull.
                    if "nicht mehr verfügbar" in body or "nicht mehr online" in body:
                        log(f"  karriere.at: '{job['title'][:40]}' expired — skipping.")
                    else:
                        desc = pg.evaluate("""
                        () => {
                            // Prefer the job-text body over the meta banner. A
                            // combined selector would return whichever comes first
                            // in document order (the banner), so query in priority.
                            const el = document.querySelector('.m-jobContent__jobText')
                                    || document.querySelector('.m-jobContent');
                            return el ? el.innerText.trim() : '';
                        }
                        """)
                        if desc:
                            job["description"] = desc[:12000]
                except Exception as e:
                    log(f"  karriere.at: desc error for {job['title']}: {e}")
                if single:
                    if job.get("description"):
                        kept = job
                        log("  ✓ karriere.at: got one description — keeping it for test run.")
                        break
                    if idx + 1 >= 4:
                        log("  ⚠ karriere.at: no description in first 4 cards.")
                        break
                    continue
                if not interruptible_sleep(random.uniform(1.5, 3.0)):
                    break

            # Test run: drop every karriere card except the one sample we scored.
            if single:
                state["jobs"] = [j for j in state["jobs"]
                                 if j.get("source") != "karriere" or j is kept]

        try:
            context.close()
        except Exception:
            pass
        browser.close()

    log(f"✓ karriere.at: {new_count} new jobs collected.")



def scrape_xing_jobs(config):
    """Scrape jobs from Xing using Playwright. Handles the consent/premium modal
    that appears before any job listings. Optional login for full access."""
    from playwright.sync_api import sync_playwright

    base_url = "https://www.xing.com"
    xing_email = config.get("xing_email", "")
    xing_password = config.get("xing_password", "")

    tf = config.get("time_filter", "r604800")
    days_map = {"r86400": "1", "r604800": "7", "r2592000": "30"}
    days = days_map.get(tf, "7")

    location = config.get("location", "")
    pages = config.get("pages", 5)
    prof = _profile(config)
    single = config.get("single_run", False)

    seen_ids = {j["id"] for j in state["jobs"]}
    new_count = 0

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent=(
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            ),
            viewport={"width": 1280, "height": 900},
            locale="de-DE",
        )
        page = context.new_page()

        def _dismiss_xing_modals(pg):
            """Accept Xing's consent/cookie/premium modal."""
            try:
                # Cookie consent button variants
                for sel in [
                    'button[data-testid="uc-accept-all-button"]',
                    'button[class*="accept"]',
                    'button[class*="consent"]',
                ]:
                    els = pg.query_selector_all(sel)
                    if els:
                        els[0].click()
                        pg.wait_for_timeout(1000)
                        return
                # Text-based buttons
                for text in ["Alle akzeptieren", "Accept all", "Zustimmen"]:
                    btn = pg.get_by_text(text, exact=True)
                    if btn.count():
                        btn.first.click()
                        pg.wait_for_timeout(1000)
                        return
            except Exception:
                pass

        # ── Login if credentials provided ────────────────────────────────────
        if xing_email and xing_password:
            log("Xing: logging in...")
            try:
                page.goto("https://login.xing.com", wait_until="domcontentloaded", timeout=30000)
                page.wait_for_timeout(2000)
                _dismiss_xing_modals(page)
                page.fill('input[name="email"], input[type="email"]', xing_email)
                page.wait_for_timeout(500)
                page.fill('input[name="password"], input[type="password"]', xing_password)
                page.wait_for_timeout(500)
                page.keyboard.press("Enter")
                page.wait_for_timeout(5000)
                if "xing.com" in page.url and "login" not in page.url:
                    log("✓ Xing: logged in.")
                else:
                    log(f"⚠ Xing: login uncertain (url={page.url}). Continuing anyway.")
            except Exception as e:
                log(f"⚠ Xing: login error: {e}. Continuing as guest.")
        else:
            log("Xing: scraping public listings (no credentials given).")

        for keyword in config["keywords"]:
            if state["stop_requested"]:
                break
            log(f"Xing: scraping '{keyword}'...")

            for page_num in range(pages):
                if state["stop_requested"]:
                    break

                search_url = (
                    f"{base_url}/jobs/search?"
                    f"keywords={keyword.replace(' ', '+')}"
                    f"&location={location.replace(' ', '+')}"
                    f"&sort=date"
                    f"&page={page_num + 1}"
                )
                try:
                    page.goto(search_url, wait_until="domcontentloaded", timeout=30000)
                    page.wait_for_timeout(2500)
                except Exception as e:
                    log(f"  Xing: nav error on page {page_num+1}: {e}")
                    break

                # Handle consent/premium modal
                _dismiss_xing_modals(page)
                page.wait_for_timeout(1500)

                # Try to close any remaining overlay (ESC)
                page.keyboard.press("Escape")
                page.wait_for_timeout(500)

                # Scroll to load lazy-rendered cards
                for _ in range(3):
                    page.keyboard.press("End")
                    page.wait_for_timeout(1000)
                page.keyboard.press("Home")
                page.wait_for_timeout(800)

                # Extract jobs using broad selectors
                jobs_data = page.evaluate("""
                () => {
                    const results = [];
                    // Xing uses hashed class names but consistent data attrs
                    const selectors = [
                        '[data-xds="JobCard"]',
                        'article[class*="job"]',
                        'li[class*="job"]',
                        '[class*="JobCard"]',
                        '[class*="job-card"]',
                        '[data-cy*="job"]',
                        'a[href*="/jobs/ad/"]',
                        'a[href*="/jobs/listing/"]',
                    ];
                    const seenHrefs = new Set();
                    for (const sel of selectors) {
                        const cards = Array.from(document.querySelectorAll(sel));
                        for (const card of cards) {
                            // Get the job link
                            const linkEl = (card.tagName === 'A' && card.href.includes('/jobs/'))
                                ? card
                                : card.querySelector('a[href*="/jobs/"]');
                            const href = linkEl ? linkEl.href : '';
                            if (!href || seenHrefs.has(href)) continue;
                            seenHrefs.add(href);

                            const titleEl = card.querySelector('h2, h3, [class*="title"], [class*="Title"]');
                            const companyEl = card.querySelector('[class*="company"], [class*="Company"], [class*="employer"]');
                            const locEl = card.querySelector('[class*="location"], [class*="Location"], [class*="city"]');
                            const salaryEl = card.querySelector('[class*="salary"], [class*="Salary"]');

                            const title = titleEl ? titleEl.textContent.trim() : '';
                            const company = companyEl ? companyEl.textContent.trim() : '';
                            const loc = locEl ? locEl.textContent.trim() : '';
                            const salary = salaryEl ? salaryEl.textContent.trim() : '';

                            if (title) {
                                results.push({ title, company, loc, salary, url: href });
                            }
                        }
                        if (results.length > 0) break;  // Stop at first selector that works
                    }
                    return results;
                }
                """)

                log(f"  Xing page {page_num+1}: {len(jobs_data)} cards")
                if not jobs_data:
                    # Log page body to help debug
                    snippet = page.evaluate("() => document.body.innerText.substring(0, 200)")
                    log(f"  Xing: page text snippet: {snippet[:100]}")
                    break

                for jd in jobs_data:
                    if state["stop_requested"]:
                        break
                    job_id = f"xing_{abs(hash(jd['url']))}"
                    if job_id in seen_ids or not jd["title"]:
                        continue
                    seen_ids.add(job_id)
                    job = {
                        "id": job_id,
                        "title": jd["title"],
                        "company": jd["company"],
                        "location": jd["loc"] or location,
                        "salary": jd["salary"],
                        "url": jd["url"],
                        "source": "xing",
                        "description": "",
                        "easy_apply": False,
                    }
                    state["jobs"].append(job)
                    new_count += 1

                delay = random.uniform(*prof["between"])
                if not interruptible_sleep(delay):
                    break

                # Test run: one page of cards is enough to find a sample job.
                if single:
                    break

            delay = random.uniform(*prof["between"])
            if not interruptible_sleep(delay):
                break

            # Test run: only the first keyword.
            if single:
                break

        # ── Fetch descriptions ────────────────────────────────────────────────
        xing_new = [j for j in state["jobs"]
                    if j.get("source") == "xing" and not j.get("description")]
        if xing_new:
            log(f"Xing: fetching descriptions for {len(xing_new)} jobs"
                + (" (test: stop at first)..." if single else "..."))
            kept = None
            for idx, job in enumerate(xing_new):
                if state["stop_requested"]:
                    break
                try:
                    page.goto(job["url"], wait_until="domcontentloaded", timeout=25000)
                    page.wait_for_timeout(2000)
                    _dismiss_xing_modals(page)
                    desc = page.evaluate("""
                    () => {
                        const el = document.querySelector(
                            '[class*="job-description"], [class*="JobDescription"], ' +
                            '[data-testid*="description"], section[class*="description"], ' +
                            '[class*="Description"]'
                        );
                        return el ? el.innerText.trim() : '';
                    }
                    """)
                    if desc:
                        job["description"] = desc[:12000]
                except Exception as e:
                    log(f"  Xing: desc error for {job['title']}: {e}")
                if single:
                    if job.get("description"):
                        kept = job
                        log("  ✓ Xing: got one description — keeping it for test run.")
                        break
                    if idx + 1 >= 4:
                        log("  ⚠ Xing: no description in first 4 cards.")
                        break
                    continue
                delay = random.uniform(2, 4)
                if not interruptible_sleep(delay):
                    break

            # Test run: drop every Xing card except the one sample we scored.
            if single:
                state["jobs"] = [j for j in state["jobs"]
                                 if j.get("source") != "xing" or j is kept]

        try:
            context.close()
        except Exception:
            pass
        browser.close()

    log(f"✓ Xing: {new_count} new jobs collected.")


def scrape_jobs(config):
    from playwright.sync_api import sync_playwright
    import re as _re

    all_jobs = []
    seen = set()           # dedup during scrape — prevents duplicate navigations
    job_counter = 0        # for periodic long pauses

    prof = _profile(config)
    single = config.get("single_run", False)
    single_tries = 0       # how many cards we've opened looking for a JD
    SINGLE_MAX_TRIES = 4    # checkpoint 2: try a few cards before flagging error
    single_done = False

    # Death-spiral guards. If the browser/context dies, every later navigation
    # throws after a full timeout — hundreds of those is the original 4-hour
    # hang. We bail the moment the session is fatally gone, or after too many
    # consecutive failures (a soft-ban / dead session), and save what we have.
    consecutive_fail = 0
    MAX_CONSECUTIVE_FAIL = 6
    session_dead = False

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/120.0.0.0 Safari/537.36",
            viewport={"width": 1280, "height": 800}
        )
        page = context.new_page()

# ── Login ──────────────────────────────────────────────────────────────
        log("Logging into LinkedIn...")
        page.goto("https://www.linkedin.com/login", wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(5000)
        try:
            # LinkedIn renders hidden autofill input stubs that come BEFORE the
            # real visible inputs in DOM order. `.first` without `:visible`
            # focuses a hidden input and keystrokes are silently lost.
            email_input = page.locator('input[type="email"]:visible').first
            email_input.wait_for(state="visible", timeout=15000)
            email_input.click()
            page.wait_for_timeout(300)
            page.keyboard.type(config["email"], delay=50)
            page.wait_for_timeout(500)
            pwd_input = page.locator('input[type="password"]:visible').first
            pwd_input.click()
            page.wait_for_timeout(300)
            page.keyboard.type(config["password"], delay=50)
            page.wait_for_timeout(500)
            page.keyboard.press("Enter")
            page.wait_for_timeout(random.randint(*prof["login_ms"]))
            current_url = page.url
            if "feed" in current_url or "jobs" in current_url:
                log(f"✓ Logged in — now on: {current_url}")
            else:
                page.screenshot(path="debug_login.png")
                log(f"✗ Login failed — ended up on: {current_url} — screenshot saved")
                browser.close()
                return
        except Exception as e:
            page.screenshot(path="debug_login.png")
            log(f"✗ Login failed: {e} — screenshot saved to debug_login.png")
            browser.close()
            return
        
        # ── Scrape each keyword ────────────────────────────────────────────────
        for keyword in config["keywords"]:
            if state["stop_requested"]:
                break
            log(f"Scraping: {keyword}")

            for page_num in range(0, config["pages"]):
                if state["stop_requested"]:
                    break

                search_url = (
                    f"https://www.linkedin.com/jobs/search/"
                    f"?keywords={keyword.replace(' ', '+')}"
                    f"&location={config['location']}"
                    f"&f_TPR={config['time_filter']}"
                    f"&sortBy=R&start={page_num * 25}"
                )

                try:
                    page.goto(search_url, timeout=30000, wait_until="domcontentloaded")
                    page.wait_for_timeout(random.randint(*prof["page_ms"]))
                except Exception as e:
                    if _is_fatal_pw_error(e):
                        log(f"  ✗ Browser session lost on search page — aborting scrape. ({str(e)[:80]})")
                        session_dead = True
                        break
                    log(f"  Timeout on page {page_num+1} for '{keyword}' — skipping")
                    break

                # Scroll to load lazy elements
                for _ in range(3):
                    page.keyboard.press("End")
                    page.wait_for_timeout(random.randint(1200, 2000))

                job_cards = (
                    page.query_selector_all(".scaffold-layout__list-item") or
                    page.query_selector_all(".jobs-search__results-list li")
                )
                log(f"  Page {page_num+1}: {len(job_cards)} cards found")

                if len(job_cards) == 0:
                    break

                # ── Extract card metadata ──────────────────────────────────────
                cards_data = []
                for card in job_cards:
                    title_el = (
                        card.query_selector(".job-card-list__title--link span[aria-hidden='true']") or
                        card.query_selector(".base-search-card__title")
                    )
                    company_el = (
                        card.query_selector(".artdeco-entity-lockup__subtitle span") or
                        card.query_selector(".base-search-card__subtitle")
                    )
                    location_el = (
                        card.query_selector(".artdeco-entity-lockup__caption li span") or
                        card.query_selector(".job-search-card__location")
                    )
                    link_el = (
                        card.query_selector("a.job-card-list__title--link") or
                        card.query_selector("a.base-card__full-link")
                    )
                    id_el = (
                        card.query_selector("[data-job-id]") or
                        card.query_selector("[data-entity-urn]")
                    )

                    if not (title_el and link_el):
                        continue

                    job_id = ""
                    if id_el:
                        job_id = id_el.get_attribute("data-job-id") or ""
                        if not job_id:
                            urn = id_el.get_attribute("data-entity-urn") or ""
                            job_id = urn.split(":")[-1] if urn else ""

                    href = link_el.get_attribute("href") or ""
                    url = href if href.startswith("http") else "https://www.linkedin.com" + href

                    # Extract numeric job ID from URL as fallback
                    if not job_id:
                        match = _re.search(r'(\d{9,})', url)
                        job_id = match.group(1) if match else url

                    cards_data.append({
                        "id": job_id,
                        "title": title_el.inner_text().strip(),
                        "company": company_el.inner_text().strip() if company_el else "",
                        "location": location_el.inner_text().strip() if location_el else "",
                        "url": url,
                        "keyword": keyword,
                    })

                # ── Fetch description for each new card ────────────────────────
                for card_data in cards_data:
                    if state["stop_requested"]:
                        break

                    job_id = card_data["id"]

                    # Skip duplicates before navigating
                    if job_id in seen:
                        continue
                    seen.add(job_id)

                    job = {
                        **card_data,
                        "source": "linkedin",
                        "description": "",
                        "salary": "",
                        "scored": False,
                    }

                    # Navigate to job page
                    match = _re.search(r'(\d{9,})', card_data["url"])
                    job_url = (
                        f"https://www.linkedin.com/jobs/view/{match.group(1)}"
                        if match else card_data["url"]
                    )

                    try:
                        page.goto(job_url, timeout=30000, wait_until="domcontentloaded")
                        page.wait_for_timeout(random.randint(*prof["page_ms"]))
                        consecutive_fail = 0  # a clean navigation = session healthy
                        if single:
                            single_tries += 1

                        # Accept cookie consent if shown
                        for cookie_text in ["Accept", "Accept cookies"]:
                            try:
                                btn = page.query_selector(f"button:has-text('{cookie_text}')")
                                if btn:
                                    btn.click(force=True)
                                    page.wait_for_timeout(1000)
                                    break
                            except:
                                pass

                        # Expand full description
                        for btn_text in ["Show more", "See more"]:
                            try:
                                btn = page.query_selector(f"button:has-text('{btn_text}')")
                                if btn:
                                    btn.click(force=True)
                                    page.wait_for_timeout(800)
                                    break
                            except:
                                pass

                        job["description"] = extract_description(page)
                        salary_conf, salary_text = extract_salary(page)
                        job["salary"] = salary_text
                        job["salary_confidence"] = salary_conf

                        job_counter += 1
                        if job_counter % 10 == 0:
                            status = "✓" if job["description"] else "⚠ no desc"
                            log(f"  [{job_counter}] {job['title'][:45]} — {status}")

                    except Exception as e:
                        if _is_fatal_pw_error(e):
                            log(f"  ✗ Browser session lost — aborting scrape, saving partial. ({str(e)[:80]})")
                            session_dead = True
                            all_jobs.append(job)
                            state["jobs"] = all_jobs
                            break
                        consecutive_fail += 1
                        log(f"  Error fetching {card_data['title'][:40]}: {str(e)[:80]} "
                            f"({consecutive_fail}/{MAX_CONSECUTIVE_FAIL})")
                        if consecutive_fail >= MAX_CONSECUTIVE_FAIL:
                            log(f"  ✗ {consecutive_fail} failures in a row — session likely "
                                f"dead or blocked. Aborting, saving partial.")
                            session_dead = True
                            all_jobs.append(job)
                            state["jobs"] = all_jobs
                            break

                    all_jobs.append(job)
                    state["jobs"] = all_jobs  # live update for status endpoint

                    # ── Single test run: stop at the first usable JD ───────────
                    if single:
                        if job.get("description"):
                            log(f"  ✓ Found a job description on try {single_tries} — scoring it.")
                            all_jobs[:] = [job]
                            state["jobs"] = all_jobs
                            single_done = True
                            break
                        if single_tries >= SINGLE_MAX_TRIES:
                            log(f"  ✗ No description found after {single_tries} cards.")
                            break
                        # try the next card quickly, no long pause
                        interruptible_sleep(random.uniform(1, 2))
                        continue

                    # ── Random wait between jobs ───────────────────────────────
                    every = prof["batch_every"]
                    if every and job_counter > 0 and job_counter % every == 0:
                        pause = random.randint(*prof["batch_pause"])
                        log(f"  ⏸ Pause {pause}s every {every} jobs...")
                        if not interruptible_sleep(pause):
                            break
                    else:
                        if not interruptible_sleep(random.uniform(*prof["between"])):
                            break

                if single_done or state["stop_requested"] or session_dead:
                    break
                # Wait between search result pages
                interruptible_sleep(random.uniform(*prof["between"]))

            if single_done or state["stop_requested"] or session_dead:
                break

        browser.close()

    log(f"Total unique jobs scraped: {len(all_jobs)}")
    log(f"  With descriptions: {len([j for j in all_jobs if j.get('description')])}")
    state["jobs"] = all_jobs

# ─── Step 2: Score Jobs ───────────────────────────────────────────────────────
def _score_one_job(config, job):
    """Score a single job in place. Safe to call from worker threads."""
    resume = config.get("resume_text", "No resume provided")

    # Only pass salary if we got it from a reliable source
    salary_line = ""
    if job.get("salary_confidence") == "verified" and job.get("salary"):
        salary_line = f"SALARY: {job['salary']}\n"

    prompt = f"""You are a hiring advisor evaluating a candidate for a specific job. Be direct — no flattery, no filler.

Compare the resume against the job requirements. Consider: seniority match, skill overlap, location/visa fit, realistic callback odds. Only mention compensation if it is listed and relevant.

RESUME:
{resume}

JOB:
Title: {job['title']}
Company: {job['company']}
Location: {job['location']}
{salary_line}
DESCRIPTION:
{job.get('description', '')[:12000]}

Return a JSON object with exactly these keys. Do NOT wrap the response in markdown or add any other text:
{{
  "fit_score": <0-100>,
  "matched_skills": ["skill1", "skill2"],
  "missing_skills": ["skill1", "skill2"],
  "response_probability": <0-100>,
  "general_gaps": "2 honest sentences on the biggest gaps",
  "resume_suggestions": ["suggestion 1", "suggestion 2", "suggestion 3"],
  "verdict": "One sentence: apply or skip, and why"
}}"""

    try:
        # 64K requested budget — clamped per-model by _resolve_max_tokens, so a
        # 64K/65K model gets full headroom and a smaller model gets its own max.
        # Scoring JSON is ~3K tokens, so this is safety headroom, not extra cost.
        result = _llm_json(config, prompt, max_tokens=64000)
        job.update(result)
        job["scored"] = True
    except Exception as e:
        log(f"  Error scoring {job['title'][:40]}: {e}")
    return job

def score_jobs(config):
    from concurrent.futures import ThreadPoolExecutor

    jobs = state["jobs"]
    to_score = [j for j in jobs if j.get("description")]
    if not to_score:
        state["scored_jobs"] = []
        log("Scoring complete: 0 jobs had descriptions to score")
        return

    # Score in parallel — the LLM call is the bottleneck, so concurrency cuts
    # wall-clock time hugely. Worker count is chosen automatically per provider
    # to stay under rate limits; the user never has to think about it. These
    # workers only make HTTP/LLM calls — they never touch Playwright.
    workers = _auto_workers(config, len(to_score))

    scored = []
    lock = threading.Lock()
    done = 0
    total = len(to_score)
    log(f"Scoring {total} jobs with {workers} parallel workers...")

    with ThreadPoolExecutor(max_workers=workers) as ex:
        futures = [ex.submit(_score_one_job, config, j) for j in to_score]
        for fut in futures:
            if state["stop_requested"]:
                fut.cancel()  # cancels only not-yet-started tasks
        for fut in futures:
            if fut.cancelled():
                continue
            try:
                job = fut.result()
            except Exception as e:
                log(f"  Worker error: {e}")
                continue
            with lock:
                scored.append(job)
                state["scored_jobs"] = list(scored)  # live update for status
                done += 1
                if done % 5 == 0 or done == total:
                    log(f"  Scored {done}/{total}")
            if state["stop_requested"]:
                log(f"  Stop requested — halting scoring at {done}/{total}")
                for f in futures:
                    f.cancel()

    state["scored_jobs"] = scored
    log(f"Scoring complete: {len([j for j in scored if j.get('scored')])} jobs scored")

# ─── Step 3: Clusters + Plan ──────────────────────────────────────────────────
def generate_clusters_and_plan(config):
    profession = config.get("profession", "the relevant field")
    jobs = state["scored_jobs"] or state["jobs"]
    resume = config.get("resume_text", "")

    all_missing = []
    for job in jobs:
        all_missing.extend(job.get("missing_skills", []))
    freq = Counter(all_missing).most_common(30)
    skill_list = "\n".join([f"- {s} ({c}x)" for s, c in freq])
    job_summaries = "\n".join([
        f"- {j['title']} @ {j['company']} | fit: {j.get('fit_score', '?')} | missing: {', '.join(j.get('missing_skills', [])[:3])}"
        for j in sorted(jobs, key=lambda x: x.get('fit_score', 0), reverse=True)[:30]
    ])

    location = config.get("location", "")
    location_line = f" targeting jobs in/around {location}." if location else ""

    prompt = f"""You are a career strategist building a personalised skill-up plan.

A candidate is applying to {len(jobs)} job listings{location_line} Their resume and aggregated skill gaps are below.

1. SKILL CLUSTERS: Group missing skills into 5-7 thematic clusters. For each: name, skills (list), score_boost (0-15 — how many fit-score points this adds), jobs_impacted (count), days_to_learn, one free resource (URL or book title), priority (high/medium/low).

2. STUDY PLAN: 4-week day-by-day learning roadmap. Highest-ROI skills first. Week 4 reserved for applications + interview prep. Each day: focus area, 2-3 concrete tasks, one deliverable, 2-4 hours.

RESUME:
{resume}

TOP JOBS (by fit score):
{job_summaries}

MISSING SKILLS (by frequency across all jobs):
{skill_list}

Return a JSON object. Do NOT wrap the response in markdown or add any other text:
{{"clusters": [{{"name":"","skills":[],"score_boost":0,"jobs_impacted":0,"days_to_learn":0,"resource":"","priority":"high"}}],
"study_plan": [{{"week":1,"theme":"","days":[{{"day":1,"focus":"","tasks":["","",""],"deliverable":"","hours":3}}]}}]}}"""

    try:
        result = _llm_json(config, prompt, max_tokens=64000)
        state["clusters"] = result.get("clusters", [])
        state["study_plan"] = result.get("study_plan", [])
        log("Skill clusters and study plan generated")
    except Exception as e:
        log(f"Error generating clusters/plan: {e}")
        state["clusters"] = []
        state["study_plan"] = []

# ─── Salary Classification ────────────────────────────────────────────────────
def classify_salary(salary_str):
    if not salary_str or not salary_str.strip():
        return "missing"
    s = salary_str.lower()
    if any(x in s for x in ["/hr", "per hour", "hourly", "/h"]):
        return "hourly"
    if any(x in s for x in ["/yr", "per year", "annually", "year", "annual", "k/y"]):
        return "annual"
    numbers = re.findall(r"[\d]+", salary_str.replace(",", ""))
    if numbers:
        try:
            val = int(numbers[0])
            if val > 1000: return "annual"
            elif val < 500: return "hourly"
        except: pass
    return "annual"

# ─── File Generators ──────────────────────────────────────────────────────────
def generate_csv():
    jobs = state["scored_jobs"] or state["jobs"]
    for job in jobs:
        job["salary_type"] = classify_salary(job.get("salary", ""))

    annual  = [j for j in jobs if j["salary_type"] == "annual"]
    hourly  = [j for j in jobs if j["salary_type"] == "hourly"]
    missing = [j for j in jobs if j["salary_type"] == "missing"]

    from datetime import datetime
    date_str = datetime.now().strftime("%Y-%m-%d")
    path = OUTPUTS_DIR / f"job_results_{date_str}.csv"

    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Source","Title","Company","Location","Salary","Salary Type",
                         "Fit Score","Response Probability","Missing Skills","Verdict","Apply Type","URL"])

        def write_section(section_jobs, label):
            if section_jobs:
                writer.writerow([f"--- {label} ({len(section_jobs)} jobs) ---"])
                for job in sorted(section_jobs, key=lambda x: x.get("fit_score", 0), reverse=True):
                    writer.writerow([
                        job.get("source",""),
                        job.get("title",""), job.get("company",""), job.get("location",""),
                        job.get("salary",""), job.get("salary_type",""),
                        job.get("fit_score",""), job.get("response_probability",""),
                        " | ".join(job.get("missing_skills",[])), job.get("verdict",""),
                        job.get("apply_type",""), job.get("url","")
                    ])
                writer.writerow([])

        write_section(annual, "ANNUAL SALARY")
        write_section(hourly, "HOURLY RATE")
        write_section(missing, "SALARY NOT LISTED")

    return str(path)

def generate_skills_txt():
    clusters = state.get("clusters", [])
    path = OUTPUTS_DIR / "skill_clusters.txt"
    with open(path, "w", encoding="utf-8") as f:
        f.write("SKILL GAP ANALYSIS\n" + "=" * 60 + "\n\n")
        if not clusters:
            jobs = state["scored_jobs"] or state["jobs"]
            all_missing = []
            for job in jobs:
                all_missing.extend(job.get("missing_skills", []))
            for skill, count in Counter(all_missing).most_common():
                f.write(f"  {count:3d}x  {skill}\n")
        else:
            for c in sorted(clusters, key=lambda x: x.get("score_boost", 0), reverse=True):
                f.write(f"[{c.get('priority','').upper()}] {c['name']}\n")
                f.write(f"  Score boost:   +{c.get('score_boost', 0)} pts\n")
                f.write(f"  Jobs impacted: {c.get('jobs_impacted', '?')}\n")
                f.write(f"  Days to learn: {c.get('days_to_learn', c.get('days', '?'))}\n")
                f.write(f"  Resource:      {c.get('resource', '')}\n")
                f.write(f"  Skills:        {', '.join(c.get('skills', []))}\n\n")
    return str(path)

def generate_plan_txt():
    plan = state.get("study_plan", [])
    path = OUTPUTS_DIR / "study_plan.txt"
    with open(path, "w", encoding="utf-8") as f:
        f.write("4-WEEK STUDY PLAN\n" + "=" * 60 + "\n\n")
        if not plan:
            f.write("Run in 'Full Analysis' mode to generate a personalized study plan.\n")
        else:
            for week in plan:
                f.write(f"WEEK {week['week']} — {week.get('theme', '')}\n" + "-" * 40 + "\n")
                for day in week.get("days", []):
                    f.write(f"\nDay {day['day']}: {day['focus']} ({day.get('hours', 3)} hrs)\n")
                    for task in day.get("tasks", []):
                        f.write(f"  □ {task}\n")
                    f.write(f"  → Deliverable: {day.get('deliverable', '')}\n")
                f.write("\n")
    return str(path)

# ─── Applier Endpoints and Functions ──────────────────────────────────────────
@app.route("/applier")
def applier():
    with open(BASE_DIR / "applier.html", encoding="utf-8") as f:
        return f.read()

@app.route("/run_applier", methods=["POST"])
def run_applier():
    global pipeline_thread

    if state["running"] or (pipeline_thread and pipeline_thread.is_alive()):
        state["stop_requested"] = True
        if pipeline_thread and pipeline_thread.is_alive():
            pipeline_thread.join(timeout=12)
        state["running"] = False

    data = request.form
    files = request.files

    config = {
        "email": data.get("email", ""),
        "password": data.get("password", ""),
        "auto_apply_min_score": int(data.get("min_score", 0) or 0),
        "auto_apply_max_score": int(data.get("max_score", 100) or 100),
        "phone_number": data.get("phone_number", ""),
        "resume_path": "",
        "llm_provider": data.get("llm_provider", "custom"),
        "llm_model": data.get("llm_model", "deepseek-chat"),
        "llm_api_key": data.get("llm_api_key", ""),
        "llm_base_url": data.get("llm_base_url", "")
    }

    # Check if they uploaded a resume or use an existing one
    if "resume" in files and files["resume"].filename:
        resume_file = files["resume"]
        suffix = Path(resume_file.filename).suffix.lower()
        attachments_dir = (PERSONAL_DIR if IS_PERSONAL else BASE_DIR) / "attachments"
        attachments_dir.mkdir(exist_ok=True)
        resume_path = attachments_dir / f"resume{suffix}"
        resume_file.save(str(resume_path))
        config["resume_path"] = str(resume_path)
    else:
        attachments_dir = (PERSONAL_DIR if IS_PERSONAL else BASE_DIR) / "attachments"
        existing = list(attachments_dir.glob("resume.*"))
        if existing:
            config["resume_path"] = str(existing[0])

    state["running"] = True
    state["stop_requested"] = False
    state["log"] = []
    state["error"] = None
    state["stage"] = "Starting Applier..."

    def _run_wrapper(cfg):
        try:
            auto_apply_jobs(cfg)
        except Exception as e:
            state["error"] = str(e)
            log(f"Error in applier: {e}")
        finally:
            state["stage"] = "Done ✓"
            state["running"] = False

    thread = threading.Thread(target=_run_wrapper, args=(config,))
    thread.daemon = True
    thread.start()
    pipeline_thread = thread

    return jsonify({"ok": True})

def extract_resume_text(path_str):
    if not path_str:
        return ""
    path = Path(path_str)
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        elif suffix in [".docx", ".doc"]:
            import docx
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        log(f"Warning: could not extract resume text: {e}")
    return ""

def _llm_json(config, prompt, max_tokens=2000, tries=2):
    """Call the LLM and parse a JSON object out of the reply, robustly.

    Layered defence against the occasional malformed reply:
      1. json_mode (response_format) makes the model emit valid JSON at source.
      2. Markdown-fence stripping + a brace-salvage regex handle prose wrapping.
      3. One retry re-rolls the rare glitch that slips through 1 and 2.
    """
    last = "no response"
    for _ in range(max(1, tries)):
        text = _llm_chat(config, prompt, max_tokens=max_tokens, json_mode=True)
        text_clean = text.strip()
        if not text_clean:
            last = "empty response"
            continue
        if text_clean.startswith("```"):
            parts = text_clean.split("```")
            if len(parts) >= 2:
                text_clean = parts[1]
            if text_clean.startswith("json"):
                text_clean = text_clean[4:]
        text_clean = text_clean.strip()
        try:
            return json.loads(text_clean)
        except Exception as e:
            last = e
            match = re.search(r'\{.*\}', text_clean, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except Exception as e2:
                    last = e2
    raise ValueError(f"Failed to parse LLM response as JSON after {tries} tries: {last}")

def linkedin_login(page, config, prof):
    log("Checking LinkedIn session...")
    try:
        page.goto("https://www.linkedin.com/feed", wait_until="domcontentloaded", timeout=30000)
        page.wait_for_timeout(3000)
        if ("feed" in page.url or "jobs" in page.url) and "login" not in page.url:
            log(f"✓ Already logged in — now on: {page.url}")
            return True
    except Exception as e:
        log(f"Session check error: {e}")
        
    log("Logging into LinkedIn...")
    page.goto("https://www.linkedin.com/login", wait_until="domcontentloaded", timeout=60000)
    page.wait_for_timeout(5000)
    try:
        email_input = page.locator('input[type="email"]:visible').first
        email_input.wait_for(state="visible", timeout=15000)
        email_input.click()
        page.wait_for_timeout(300)
        page.keyboard.type(config["email"], delay=50)
        page.wait_for_timeout(500)
        pwd_input = page.locator('input[type="password"]:visible').first
        pwd_input.click()
        page.wait_for_timeout(300)
        page.keyboard.type(config["password"], delay=50)
        page.wait_for_timeout(500)
        page.keyboard.press("Enter")
        page.wait_for_timeout(random.randint(*prof["login_ms"]))
        current_url = page.url
        if "checkpoint" in current_url or "challenge" in current_url:
            log("[!] LinkedIn security check/captcha encountered! Please solve it in the visible browser window...")
            for attempt in range(300):
                page.wait_for_timeout(1000)
                current_url = page.url
                if "feed" in current_url or "jobs" in current_url:
                    log("✓ Captcha solved! Continuing...")
                    break
            else:
                log("✗ Captcha not solved within 300 seconds. Aborting.")
                return False
                
        if "feed" in current_url or "jobs" in current_url:
            log(f"✓ Logged in — now on: {current_url}")
            return True
        else:
            page.screenshot(path="debug_login.png")
            log(f"✗ Login failed — ended up on: {current_url} — screenshot saved")
            return False
    except Exception as e:
        page.screenshot(path="debug_login.png")
        log(f"✗ Login failed: {e} — screenshot saved to debug_login.png")
        return False

def auto_apply_jobs(config):
    from playwright.sync_api import sync_playwright

    if not config.get("resume_text") and config.get("resume_path"):
        log("Extracting resume text from file...")
        config["resume_text"] = extract_resume_text(config["resume_path"])

    min_score = config.get("auto_apply_min_score", 0)
    max_score = config.get("auto_apply_max_score", 100)
    eligible_jobs = [j for j in state["scored_jobs"] if j.get("is_easy_apply") and min_score <= int(j.get("fit_score", 0) if j.get("fit_score", 0) != "" else 0) <= max_score]
    
    if not eligible_jobs:
        log("No jobs met the criteria for Auto-Apply (Easy Apply + score threshold).")
        return

    log(f"Starting Auto-Apply for {len(eligible_jobs)} job(s)...")
    prof = _profile(config)

    # JS snippet to extract fields
    EXTRACT_JS = """
    () => {
        const dialog = document.querySelector('dialog');
        if (!dialog) return [];
        
        const fields = [];
        dialog.querySelectorAll('input, select, textarea').forEach(el => {
            const type = el.type || el.tagName.toLowerCase();
            if (type === 'file' || type === 'hidden') return;
            
            let label = '';
            if (el.id) {
                const labelEl = dialog.querySelector(`label[for="${el.id}"]`);
                if (labelEl) label = labelEl.textContent;
            }
            
            if (!label) {
                let parent = el.parentElement;
                let depth = 0;
                while (parent && parent !== dialog && depth < 5) {
                    const labelEl = parent.querySelector('label, legend, .fb-form-element__label, .fb-text-textarea__label-title, .fb-text-textarea__label, [class*="label"], [class*="title"], h3, h4, h5');
                    if (labelEl && labelEl.textContent.trim()) {
                        label = labelEl.textContent;
                        break;
                    }
                    parent = parent.parentElement;
                    depth++;
                }
            }
            
            if (!label) {
                let parent = el.parentElement;
                let depth = 0;
                while (parent && parent !== dialog && depth < 5) {
                    const pEl = parent.querySelector('p');
                    if (pEl && pEl.textContent.trim()) {
                        const txt = pEl.textContent.trim();
                        if (!txt.match(/^\d+\/\d+/) && !txt.match(/^\d+ of \d+/) && !txt.match(/^0 of \d+/)) {
                            label = txt;
                            break;
                        }
                    }
                    parent = parent.parentElement;
                    depth++;
                }
            }
            
            label = label ? label.trim().replace(/\\n/g, ' ') : '';
            
            const nameOrId = el.name || el.id;
            if (!nameOrId) return;
            if (fields.some(f => f.id === nameOrId)) return;
            
            let isFilled = false;
            if (type === 'text' || type === 'textarea' || type === 'email' || type === 'tel' || type === 'number') {
                isFilled = el.value.trim() !== '';
            } else if (type === 'select-one' || el.tagName.toLowerCase() === 'select') {
                const val = (el.value || '').trim().toLowerCase();
                const text = el.options && el.selectedIndex >= 0 ? el.options[el.selectedIndex].text.trim().toLowerCase() : '';
                const isPlaceholder = val === '' ||
                                      val.includes('select') ||
                                      val.includes('choose') ||
                                      val.includes('please') ||
                                      text.includes('select') ||
                                      text.includes('choose') ||
                                      text.includes('please') ||
                                      val === '-' ||
                                      val === '--';
                isFilled = !isPlaceholder;
            } else if (type === 'radio' || type === 'checkbox') {
                const groupInputs = el.name ? Array.from(dialog.querySelectorAll(`input[name="${el.name}"]`)) : [el];
                isFilled = groupInputs.some(i => i.checked);
            }
            
            if (isFilled) return;
            
            let options = [];
            if (el.tagName.toLowerCase() === 'select') {
                options = Array.from(el.querySelectorAll('option'))
                    .map(o => o.textContent.trim())
                    .filter(v => v && v !== 'Select an option');
            } else if (type === 'radio') {
                const groupInputs = el.name ? Array.from(dialog.querySelectorAll(`input[name="${el.name}"]`)) : [el];
                groupInputs.forEach(r => {
                    let rLabelText = '';
                    if (r.id) {
                        const rLabel = dialog.querySelector(`label[for="${r.id}"]`);
                        if (rLabel) rLabelText = rLabel.textContent.trim();
                    }
                    if (!rLabelText) {
                        const parent = r.parentElement;
                        if (parent) rLabelText = parent.textContent.trim();
                    }
                    options.push(r.value + (rLabelText ? " (" + rLabelText + ")" : ""));
                });
            }
            
            fields.push({
                id: nameOrId,
                label: label,
                type: type,
                options: options
            });
        });
        return fields;
    }
    """

    with sync_playwright() as p:
        user_data_dir = os.path.join(os.path.expanduser("~"), ".linkedin_playwright_profile")
        try:
            context = p.chromium.launch_persistent_context(
                user_data_dir,
                headless=False,
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                viewport={"width": 1280, "height": 800}
            )
            page = context.pages[0] if context.pages else context.new_page()
            if not linkedin_login(page, config, prof):
                log("✗ LinkedIn login failed — aborting Auto-Apply.")
                return

            applied_data = load_file(APPLIED_FILE)

            for idx, job in enumerate(eligible_jobs):
                if state["stop_requested"]:
                    break
                
                log(f"  Attempting Auto-Apply [{idx+1}/{len(eligible_jobs)}]: {job['title'][:40]}")
                
                try:
                    page.goto(job["url"], timeout=30000, wait_until="domcontentloaded")
                    page.wait_for_timeout(random.randint(2000, 4000))
                    
                    apply_btn = page.query_selector(".jobs-apply-button, button[aria-label*='Easy Apply'], button[aria-label*='Einfach bewerben']")
                    if not apply_btn:
                        log(f"    ✗ Easy Apply button not found on page.")
                        state["applier_results"].append({
                            "title": job.get("title", ""),
                            "company": job.get("company", ""),
                            "url": job.get("url", ""),
                            "status": "Failed",
                            "error": "Easy Apply button not found"
                        })
                        continue
                        
                    apply_btn.click(force=True)
                    page.wait_for_timeout(1500)
                    
                    max_steps = 10
                    step_count = 0
                    submitted = False
                    
                    while step_count < max_steps and not state["stop_requested"]:
                        step_count += 1
                        page.wait_for_timeout(1000)
                        
                        # Handle Phone Number specifically if we can, to save LLM tokens
                        phone_input = page.query_selector("input[id*='phoneNumber']")
                        if phone_input and phone_input.is_visible() and config.get("phone_number"):
                            phone_input.fill(config["phone_number"])
                            
                        # Handle Resume Upload
                        file_input = page.query_selector("input[type='file']")
                        if file_input and config.get("resume_path") and Path(config["resume_path"]).exists():
                            file_input.set_input_files(config["resume_path"])
                            page.wait_for_timeout(1500)

                        # Extract remaining empty fields
                        empty_fields = page.evaluate(EXTRACT_JS)
                        
                        if empty_fields:
                            log(f"    [+] Found {len(empty_fields)} empty field(s). Asking LLM for answers...")
                            fields_summary = ", ".join([f"{f.get('label') or f.get('id')} ({f.get('type')})" for f in empty_fields[:5]])
                            if len(empty_fields) > 5:
                                fields_summary += f" and {len(empty_fields)-5} more..."
                            log(f"      Fields: {fields_summary}")
                            
                            prompt = (
                                f"You are an AI assistant helping a user apply for a job on LinkedIn.\n"
                                f"Job Title: {job.get('title')}\n"
                                f"Company: {job.get('company')}\n"
                                f"Job Description:\n{job.get('description', '')[:3000]}\n\n"
                                f"User Resume:\n{config.get('resume_text', '')[:5000]}\n\n"
                                f"The application form requires answers to the following fields. "
                                f"Return a strict JSON object where the keys are the field 'id's and the values are the appropriate answers.\n"
                                f"If it is a multiple choice/select/radio field, you MUST return one of the EXACT options provided.\n"
                                f"If you do not know the answer, infer the most professional and likely true answer based on the resume (e.g. 0 for years of experience with a tool not on the resume). If asked about sponsorship/visa, default to 'No' unless the resume implies otherwise.\n"
                                f"CRITICAL: For checkboxes/consents representing terms, privacy policy, agreement, accuracy, data processing, or acknowledgement, you MUST answer true / yes / agree to prevent application validation failure.\n\n"
                                f"Fields:\n{empty_fields}\n\n"
                                f"Respond ONLY with the JSON dictionary. Do not include markdown blocks or any other text."
                            )
                            
                            try:
                                answers = _llm_json(config, prompt, max_tokens=4000)
                                log(f"      LLM answered: {json.dumps(answers)}")
                                for field in empty_fields:
                                    # Normalize key lookup to prevent guillemet mismatch
                                    ans = answers.get(field["id"])
                                    if ans is None:
                                        # Try stripping guillemets
                                        clean_id = field["id"].replace("«", "").replace("»", "").strip()
                                        ans = answers.get(clean_id)
                                        if ans is None:
                                            # Try matching clean keys in answers dict
                                            for k, v in answers.items():
                                                if k.replace("«", "").replace("»", "").strip() == clean_id:
                                                    ans = v
                                                    break
                                    
                                    if ans is not None:
                                        # Use playwright to fill
                                        f_id = field["id"]
                                        f_type = field["type"]
                                        log(f"      -> Filling {f_id} ({field['label']}) with: {ans}")
                                        if f_type in ["text", "textarea", "email", "tel", "number"]:
                                            inp = page.query_selector(f"[id='{f_id}'], [name='{f_id}']")
                                            if inp:
                                                inp.fill(str(ans))
                                                page.wait_for_timeout(200)
                                                
                                                # Autocomplete helper for locations, schools, companies
                                                lbl_lower = field.get("label", "").lower()
                                                is_autocomplete = any(x in lbl_lower for x in [
                                                    "location", "city", "ort", "stadt", "plz", "country", "land", "state",
                                                    "company", "employer", "arbeitgeber", "firma", "school", "university",
                                                    "schule", "universität", "hochschule"
                                                ])
                                                if is_autocomplete:
                                                    page.wait_for_timeout(1000) # wait for autocomplete dropdown
                                                    inp.press("ArrowDown")
                                                    page.wait_for_timeout(400)
                                                    inp.press("Enter")
                                                    page.wait_for_timeout(400)
                                                    
                                                inp.press("Tab")
                                                page.wait_for_timeout(200)
                                        elif f_type in ["select-one", "select"]:
                                            sel = page.query_selector(f"[id='{f_id}'], [name='{f_id}']")
                                            if sel:
                                                ans_str = str(ans).strip()
                                                try:
                                                    sel.select_option(value=ans_str)
                                                except Exception:
                                                    try:
                                                        sel.select_option(label=ans_str)
                                                    except Exception:
                                                        try:
                                                            options = sel.evaluate("el => Array.from(el.options).map(o => ({value: o.value, text: o.text}))")
                                                            matched_value = None
                                                            for opt in options:
                                                                if ans_str.lower() in opt["text"].lower() or ans_str.lower() in opt["value"].lower():
                                                                    matched_value = opt["value"]
                                                                    break
                                                            if matched_value:
                                                                sel.select_option(value=matched_value)
                                                            else:
                                                                log(f"        [!] Could not match select option for: {ans_str}")
                                                        except Exception as sel_err:
                                                            log(f"        [!] Failed to select option: {sel_err}")
                                        elif f_type == "radio":
                                            rads = page.query_selector_all(f"input[type='radio'][name='{f_id}']")
                                            if not rads:
                                                rads = page.query_selector_all(f"input[type='radio'][name*='{f_id}']")
                                            
                                            if rads:
                                                ans_str = str(ans).lower().strip()
                                                selected_rad = None
                                                
                                                # 1. Evaluate values and labels for all radios in the group
                                                radio_options = []
                                                for r in rads:
                                                     val = r.evaluate("el => el.value").lower().strip()
                                                     r_id = r.get_attribute("id")
                                                     lbl_text = ""
                                                     if r_id:
                                                         lbl = page.query_selector(f"label[for='{r_id}']")
                                                         if lbl:
                                                             lbl_text = lbl.text_content().lower().strip()
                                                     if not lbl_text:
                                                         lbl_text = r.evaluate("el => el.parentElement ? el.parentElement.textContent : ''").lower().strip()
                                                     radio_options.append({"el": r, "value": val, "label": lbl_text})
                                                     
                                                # 2. Try exact match on value
                                                for opt in radio_options:
                                                     if opt["value"] == ans_str:
                                                         selected_rad = opt["el"]
                                                         break
                                                         
                                                # 3. Try match on label text
                                                if not selected_rad:
                                                     for opt in radio_options:
                                                         if ans_str == opt["label"] or ans_str in opt["label"]:
                                                             selected_rad = opt["el"]
                                                             break
                                                             
                                                # 4. Try yes/no heuristics for 2-option groups using labels
                                                if not selected_rad and len(radio_options) == 2:
                                                     lbl0 = radio_options[0]["label"]
                                                     lbl1 = radio_options[1]["label"]
                                                     is_lbl0_yes = any(y in lbl0 for y in ["yes", "ja", "true", "y", "ok", "agree"])
                                                     is_lbl1_no = any(n in lbl1 for n in ["no", "nein", "false", "n", "disagree"])
                                                     
                                                     if is_lbl0_yes or is_lbl1_no:
                                                         if ans_str in ["yes", "y", "true", "t", "1", "ja", "on", "checked", "agree"]:
                                                             selected_rad = radio_options[0]["el"]
                                                         elif ans_str in ["no", "n", "false", "f", "0", "nein", "off", "disagree"]:
                                                             selected_rad = radio_options[1]["el"]
                                                             
                                                # 5. Fallback
                                                if not selected_rad and ans_str in ["on", "true", "yes", "ja"]:
                                                     selected_rad = radio_options[0]["el"]
                                                 
                                                if selected_rad:
                                                    try:
                                                        selected_rad.check(force=True)
                                                    except Exception:
                                                        lbl = page.query_selector(f"label[for='{selected_rad.get_attribute('id')}']") if selected_rad.get_attribute('id') else None
                                                        if lbl:
                                                            lbl.click(force=True)
                                                        else:
                                                            selected_rad.evaluate("(el) => { el.checked = true; el.dispatchEvent(new Event('change', { bubbles: true })); }")
                                                else:
                                                    log(f"        [!] Could not match radio answer '{ans}' to any option.")
                                            else:
                                                lbl = page.query_selector(f"label:has-text('{ans}')")
                                                if lbl: lbl.click(force=True)
                                        elif f_type == "checkbox":
                                            cb = page.query_selector(f"input[type='checkbox'][name='{f_id}'], input[type='checkbox'][id='{f_id}']")
                                            if cb:
                                                should_check = str(ans).lower() in ["true", "yes", "y", "1", "checked", "on"]
                                                lbl = page.query_selector(f"label[for='{cb.get_attribute('id')}']") if cb.get_attribute('id') else None
                                                if lbl:
                                                    try:
                                                        is_checked = cb.is_checked()
                                                        if is_checked != should_check:
                                                            lbl.click(force=True)
                                                            page.wait_for_timeout(100)
                                                            # Verify state updated, fallback if not
                                                            if cb.is_checked() != should_check:
                                                                cb.evaluate(f"(el) => {{ el.checked = {str(should_check).lower()}; el.dispatchEvent(new Event('change', {{ bubbles: true }})); }}")
                                                    except Exception:
                                                        try:
                                                            cb.set_checked(should_check, force=True)
                                                        except Exception:
                                                            cb.evaluate(f"(el) => {{ el.checked = {str(should_check).lower()}; el.dispatchEvent(new Event('change', {{ bubbles: true }})); }}")
                                                else:
                                                    try:
                                                        cb.set_checked(should_check, force=True)
                                                    except Exception:
                                                        try:
                                                            cb.evaluate(f"(el) => {{ el.checked = {str(should_check).lower()}; el.dispatchEvent(new Event('change', {{ bubbles: true }})); }}")
                                                        except Exception as cb_err:
                                                            log(f"        [!] Failed to set checkbox via JS: {cb_err}")
                                log(f"    [✓] Filled LLM answers.")
                            except Exception as llm_e:
                                log(f"    [!] Failed to get LLM answers: {llm_e}")
                            
                            page.wait_for_timeout(1000)
                            
                        # Check for error states (after filling attempt)
                        err_msg = page.query_selector(".artdeco-inline-feedback--error, [class*='feedback--error'], [class*='error-message'], .fb-form-element__feedback")
                        if err_msg and err_msg.is_visible():
                            log(f"    ✗ Encountered form error (missing/invalid field). Aborting this application.")
                            dismiss = page.query_selector("button[data-test-modal-close-btn], button[aria-label='Dismiss']")
                            if dismiss:
                                dismiss.click(force=True)
                                page.wait_for_timeout(1000)
                                confirm = page.query_selector("button[data-test-dialog-primary-btn], button:has-text('Discard'), button:has-text('Not now')")
                                if confirm: confirm.click(force=True)
                            state["applier_results"].append({
                                "title": job.get("title", ""),
                                "company": job.get("company", ""),
                                "url": job.get("url", ""),
                                "status": "Failed",
                                "error": "Form error (missing or invalid field)"
                            })
                            break
                            
                        # Look for Submit / Review / Next buttons
                        submit_btn = page.query_selector("button[aria-label='Submit application']")
                        review_btn = page.query_selector("button[aria-label='Review your application']")
                        next_btn = page.query_selector("button[aria-label='Continue to next step']")
                        
                        btn_to_click = None
                        is_submit = False
                        
                        if submit_btn and submit_btn.is_visible():
                            btn_to_click = submit_btn
                            is_submit = True
                        elif review_btn and review_btn.is_visible():
                            btn_to_click = review_btn
                        elif next_btn and next_btn.is_visible():
                            btn_to_click = next_btn
                        else:
                            next_fallback = page.query_selector("button:has-text('Next'), button:has-text('Weiter'), button:has-text('Continue')")
                            review_fallback = page.query_selector("button:has-text('Review')")
                            submit_fallback = page.query_selector("button:has-text('Submit')")
                            
                            if submit_fallback and submit_fallback.is_visible():
                                btn_to_click = submit_fallback
                                is_submit = True
                            elif review_fallback and review_fallback.is_visible():
                                btn_to_click = review_fallback
                            elif next_fallback and next_fallback.is_visible():
                                btn_to_click = next_fallback
                        
                        if btn_to_click:
                            # Check if disabled
                            is_disabled = (btn_to_click.get_attribute("disabled") is not None or 
                                           btn_to_click.get_attribute("aria-disabled") == "true")
                            if is_disabled:
                                log("    [!] Progress button is disabled. Form validation might have failed.")
                            
                            try:
                                btn_to_click.click(timeout=5000)
                            except Exception:
                                try:
                                    btn_to_click.click(force=True, timeout=5000)
                                except Exception as click_err:
                                    log(f"    [!] Click failed, trying JS click: {click_err}")
                                    try:
                                        btn_to_click.evaluate("(el) => el.click()")
                                    except Exception as js_err:
                                        log(f"    [!] JS click failed: {js_err}")
                            
                            if is_submit:
                                page.wait_for_timeout(2500)
                                submitted = True
                                break
                            else:
                                page.wait_for_timeout(1500)
                        else:
                            log(f"    ✗ Could not find Next/Review/Submit buttons. Aborting.")
                            break
                                
                    if submitted:
                        log(f"    ✓ Successfully Auto-Applied!")
                        remove_from_all(job["id"])
                        applied_data[job["id"]] = job
                        save_file(APPLIED_FILE, applied_data)
                    else:
                        if not state["stop_requested"] and step_count >= max_steps:
                            log(f"    ✗ Exceeded maximum modal steps. Aborting.")
                            state["applier_results"].append({
                                "title": job.get("title", ""),
                                "company": job.get("company", ""),
                                "url": job.get("url", ""),
                                "status": "Failed",
                                "error": "Exceeded maximum modal steps (Stuck)"
                            })
                except Exception as e:
                    import traceback
                    tb = traceback.format_exc()
                    log(f"    ✗ Error during auto-apply:\n{tb}")
                    state["applier_results"].append({
                        "title": job.get("title", ""),
                        "company": job.get("company", ""),
                        "url": job.get("url", ""),
                        "status": "Failed",
                        "error": str(e)
                    })

        finally:
            try:
                context.close()
            except Exception:
                pass

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    host = os.environ.get("HOST", "127.0.0.1")
    print("Starting Job Scraper Dashboard...")
    print(f"Open http://localhost:{port} in your browser")
    app.run(debug=False, host=host, port=port)